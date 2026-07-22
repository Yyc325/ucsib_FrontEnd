from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path("/Users/milescai/Downloads/Criterion C.docx")
OUTPUT = Path("/Users/milescai/ucsib45_web/Criterion C - revised v2.docx")
ASSET_DIR = Path("/Users/milescai/ucsib45_web/tmp/criterion-c-edit")


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_p.addnext(new_paragraph._p)
    new_p.getparent().remove(new_p)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def insert_table_after(document, paragraph, rows, widths):
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.autofit = False
    table_element = table._tbl
    table_element.getparent().remove(table_element)
    paragraph._p.addnext(table_element)

    for idx, text in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        cell.text = text
        cell.width = widths[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "DCE6F1")
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_data in rows[1:]:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = text
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            )

    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAB4C3")
        borders.append(border)
    return table


def add_picture_after(paragraph, path, width):
    picture_paragraph = insert_paragraph_after(paragraph)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(path), width=width)
    return picture_paragraph


def replace_text(document, old, new):
    paragraph = find_paragraph(document, old)
    paragraph.text = new
    return paragraph


document = Document(SOURCE)

# Remove narration about drawing tools and use the word budget for system explanation.
replace_text(
    document,
    "All diagrams and design assets in this document were created using two tools. Figma was used to design all UI wireframes and page layouts, allowing visual planning of component placement and user flows before frontend development. Mermaid was used to produce structural diagrams such as DFD Level 0, UI Navigation Map, the ERD and UML Class Diagram, using its text-based syntax to generate clean, consistent visuals directly from code.",
    "Figma was used for interface wireframes, while Mermaid was used for architecture, database and algorithm diagrams.",
)

caption_01 = find_paragraph(document, "Fig. 0.1")
system_description = insert_paragraph_after(
    caption_01,
    "Figure 0.1 shows a layered design: Vue handles presentation and client state, Axios sends API requests, Django applies authentication and business rules, and MySQL or cloud storage persists records and uploaded images.",
    "normal",
)

caption_03 = find_paragraph(document, "Fig. 0.3")
insert_paragraph_after(
    caption_03,
    "Figures 0.2 and 0.3 define the data exchanged with administrators and public users, then separate public pages from authenticated community actions and role-restricted dashboard functions.",
    "normal",
)

caption_17 = find_paragraph(document, "Fig. 1.7")
insert_paragraph_after(
    caption_17,
    "Figure 1.7 shows the notice publishing and backup relationships. Comment authorship is logical because the authenticated display name is stored in the comment rather than through a physical foreign key.",
    "normal",
)

# Correct module descriptions so they match the implemented UI and authentication flow.
replace_text(
    document,
    "Wireframe of GUI: Login card with phone, password fields, and role selection.",
    "Wireframe of GUI: Login and registration forms collect phone, password and account details.",
)
replace_text(
    document,
    "Function Description: Validates identity using MD5 hashing and issues a JWT token for session persistence.",
    "Function Description: Compares the MD5-derived credential and issues a JWT containing the authenticated user's identity.",
)
replace_text(
    document,
    "Function Description: An Asynchronous Thread that auto-updates notice status when the target time is reached.",
    "Function Description: A background thread changes pending notices to published when their scheduled time is reached.",
)
replace_text(
    document,
    "Function Description: Handles peer-to-peer communication using DRF Serializers to validate and save student comments. backend validates login token before saving the comment, and author is taken from the authenticated user.",
    "Function Description: DRF validates comment data; the backend verifies the JWT and stores the authenticated user's display name as author.",
)

# Scale the original search-flow image to the usable page area and remove the
# spacer paragraphs that otherwise produce an empty page before the figure.
search_caption = find_paragraph(document, "Fig. 2.5.1")
cursor = search_caption._p.getnext()
while cursor is not None:
    xml = cursor.xml
    if "image8.png" in xml:
        for extent in cursor.xpath(".//wp:extent | .//a:ext"):
            extent.set("cx", "1727000")
            extent.set("cy", "7300000")
        break
    next_cursor = cursor.getnext()
    if not cursor.xpath(".//w:t") and not cursor.xpath(".//w:drawing"):
        cursor.getparent().remove(cursor)
    cursor = next_cursor

# Complete the previously empty mobile-design section.
mobile_heading = replace_text(
    document,
    "6. Mobile Fitting(Success Criterion 6)",
    "6. Mobile Responsive Design (Success Criterion 6)",
)
mobile_heading.paragraph_format.page_break_before = True
mobile_description = insert_paragraph_after(
    mobile_heading,
    "Responsive CSS media queries adapt the interface at widths of 575.98 pixels or below. Desktop navigation is replaced by a collapsible menu, multi-column content becomes full-width or vertical, and non-essential panels are hidden to prevent horizontal scrolling.",
    "normal",
)

comparison = add_picture_after(
    mobile_description,
    ASSET_DIR / "mobile-layout-comparison.png",
    Inches(6.2),
)
comparison_caption = insert_paragraph_after(
    comparison,
    "Fig. 2.6.1 Desktop and Mobile Layout Comparison",
    "normal",
)
comparison_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

table_rows = [
    ["Element", "Desktop design", "Mobile design (<=575.98 px)"],
    ["Navigation", "Full horizontal menu and search", "Collapsible menu; desktop search hidden"],
    ["News", "Featured area and notice list", "Notice list expands to full width"],
    ["Community", "Side-by-side sections", "Sections stack and images resize"],
    ["Login", "Decorative panel and form", "Panel hidden; form uses full width"],
]
responsive_table = insert_table_after(
    document,
    comparison_caption,
    table_rows,
    [Inches(1.05), Inches(2.35), Inches(2.9)],
)

table_anchor = insert_paragraph_after(responsive_table.rows[-1].cells[-1].paragraphs[-1])
# The helper above inserts inside the last cell, so move the paragraph after the table.
table_anchor._p.getparent().remove(table_anchor._p)
responsive_table._tbl.addnext(table_anchor._p)

flow_picture = add_picture_after(
    table_anchor,
    ASSET_DIR / "mobile-responsive-flow.png",
    Inches(4.7),
)
flow_caption = insert_paragraph_after(
    flow_picture,
    "Fig. 2.6.2 Responsive Layout Decision Flow",
    "normal",
)
flow_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Keep the UML on a clean page and explain its non-obvious relationships.
uml_heading = find_paragraph(document, "3. UML Diagram")
uml_heading.paragraph_format.page_break_before = True
uml_caption = find_paragraph(document, "Fig. 3.1")
insert_paragraph_after(
    uml_caption,
    "Figure 3.1 maps request handlers to models and serializers. Notice ownership uses a database relationship, whereas comment identity is supplied from the authenticated JWT payload.",
    "normal",
)

# Criterion C contains a planned strategy; actual outcomes belong in Criterion D.
replace_text(
    document,
    "I will perform backend integration testing using Postman to simulate frontend API calls. This allows for direct verification of JSON responses, HTTP status codes, and JWT validation logic. The focus is on ensuring data integrity in the MySQL database and proper error handling for invalid inputs.",
    "Postman will simulate frontend API requests and verify JSON responses, HTTP status codes, JWT authorization and resulting MySQL changes for normal, boundary and invalid cases.",
)
replace_text(document, "Tested through Postman", "Planned for execution through Postman")
replace_text(
    document,
    "4.2.5 T11-T12. Mobile Compatibility (Related SC: #6)",
    "4.2.5 T12-T13. Mobile Compatibility (Related SC: #6)",
)

# Prevent labels and subsection headings from being stranded at page bottoms.
for paragraph in document.paragraphs:
    text = paragraph.text.strip()
    if (
        text.startswith("Table: ")
        or text == "ERD Diagram"
        or text.startswith("4.2.4 Global Search")
        or text.startswith("4.2.5 T12-T13")
    ):
        paragraph.paragraph_format.keep_with_next = True

# Correct specific test descriptions without changing the existing table structure.
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            replacements = {
                "Record saved in admin_role_notice.": "Record saved in admin_notice_copy1.",
                "with student author link.": "with the JWT-derived author name.",
                "400 Bad Requests": "400 Bad Request",
                "Middleware blocks request": "Backend authorization blocks request",
            }
            for paragraph in cell.paragraphs:
                for old, new in replacements.items():
                    if old in paragraph.text:
                        paragraph.text = paragraph.text.replace(old, new)

# T5 is the automatic-publishing test and therefore maps to SC3.
for table in document.tables:
    for row in table.rows:
        row_text = " ".join(cell.text for cell in row.cells)
        if row_text.strip().startswith("T5 "):
            row.cells[1].text = "#3"

# Use consistent spacing for the newly added explanation paragraphs.
for paragraph in document.paragraphs:
    if paragraph.text.startswith("Figure ") or paragraph.text.startswith("Figures "):
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(8)

document.save(OUTPUT)
print(OUTPUT)
